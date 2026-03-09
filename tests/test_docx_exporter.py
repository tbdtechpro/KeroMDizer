"""Tests for the DOCX exporter."""
import pytest
from pathlib import Path
from docx import Document
from docx_exporter import export_docx


MINIMAL_MD = """\
# My Conversation

_2026-01-14  ·  gpt-4o_

---

### \U0001f464 User

Hello world

---

### \U0001f916 Assistant

Here is some code:
```python
print('hello')
```

---
"""


def test_docx_file_created(tmp_path):
    out = tmp_path / 'test.docx'
    export_docx(MINIMAL_MD, out)
    assert out.exists()
    assert out.stat().st_size > 0


def test_docx_title_heading(tmp_path):
    out = tmp_path / 'test.docx'
    export_docx(MINIMAL_MD, out)
    doc = Document(str(out))
    headings = [p for p in doc.paragraphs if getattr(p.style, 'name', '').startswith('Heading 1')]
    assert any('My Conversation' in p.text for p in headings)


def test_docx_role_headers(tmp_path):
    out = tmp_path / 'test.docx'
    export_docx(MINIMAL_MD, out)
    doc = Document(str(out))
    headings = [p for p in doc.paragraphs if getattr(p.style, 'name', '').startswith('Heading 3')]
    texts = [p.text for p in headings]
    assert any('User' in t for t in texts)
    assert any('Assistant' in t for t in texts)


def test_docx_code_block_monospace(tmp_path):
    out = tmp_path / 'test.docx'
    export_docx(MINIMAL_MD, out)
    doc = Document(str(out))
    courier_runs = [
        run
        for para in doc.paragraphs
        for run in para.runs
        if run.font.name == 'Courier New'
    ]
    assert len(courier_runs) > 0


def test_docx_creates_parent_dirs(tmp_path):
    out = tmp_path / 'nested' / 'dir' / 'test.docx'
    export_docx(MINIMAL_MD, out)
    assert out.exists()


def test_docx_prose_present(tmp_path):
    out = tmp_path / 'test.docx'
    export_docx(MINIMAL_MD, out)
    doc = Document(str(out))
    all_text = ' '.join(p.text for p in doc.paragraphs)
    assert 'Hello world' in all_text
