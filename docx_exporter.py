"""DOCX exporter for KeroMDizer conversations."""
from __future__ import annotations
import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


_FENCE_RE = re.compile(r'```(\w*)\n(.*?)```', re.DOTALL)


def _add_code_block(doc: DocumentType, code: str, language: str = '') -> None:
    """Add a code block paragraph with monospace styling and grey background."""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']  # type: ignore[assignment]
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F6F8FA')
    pPr.append(shd)


def _add_prose(doc: DocumentType, text: str) -> None:
    """Add prose text as a paragraph, skipping empty blocks."""
    text = text.strip()
    if text:
        doc.add_paragraph(text)


def export_docx(md_content: str, output_path: Path) -> None:
    """Convert rendered markdown string to a .docx file."""
    doc = Document()

    sections = re.split(r'\n---\n', md_content)

    for section in sections:
        section = section.strip()
        if not section:
            continue

        # H1 title
        first_line = section.split('\n', 1)[0]
        if first_line.startswith('# '):
            doc.add_heading(first_line[2:].strip(), level=1)
            rest = section[len(first_line):].strip()
            if rest:
                _add_prose(doc, rest)
            continue

        # Italic metadata line: _date · model_
        if section.startswith('_') and section.endswith('_') and '\n' not in section:
            p = doc.add_paragraph()
            run = p.add_run(section[1:-1])
            run.italic = True
            run.font.color.rgb = RGBColor(0x6A, 0x73, 0x7D)
            continue

        # Role header + body: ### 👤 User  or  ### 🤖 Assistant
        if section.startswith('### '):
            header_line, _, body = section.partition('\n')
            heading = doc.add_heading(header_line[4:].strip(), level=3)
            if heading.runs:
                if '👤' in header_line or 'User' in header_line:
                    heading.runs[0].font.color.rgb = RGBColor(0x03, 0x66, 0xD6)
                else:
                    heading.runs[0].font.color.rgb = RGBColor(0x28, 0xA7, 0x45)

            # Parse body: alternate prose and code fences
            parts = _FENCE_RE.split(body)
            i = 0
            while i < len(parts):
                if i % 3 == 0:
                    _add_prose(doc, parts[i])
                else:
                    lang = parts[i]
                    code = parts[i + 1] if i + 1 < len(parts) else ''
                    _add_code_block(doc, code.strip(), lang)
                    i += 1
                i += 1
            continue

        # Fallback: treat as prose
        _add_prose(doc, section)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
